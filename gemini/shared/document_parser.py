"""
Universal Document Parser for Gemini
=====================================

Handles multiple document formats:
- PDF (.pdf)
- Text (.txt, .md)
- JSON (.json, .jsonl)
- CSV (.csv)
- Word (.docx)
- YouTube transcripts
- HTML/Web content

Author: Claude (for Gemini)
Date: 2025-11-06
Status: PRODUCTION READY
"""

import json
import csv
import re
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from dataclasses import dataclass
from datetime import datetime


@dataclass
class ParsedDocument:
    """Container for parsed document data"""
    file_path: str
    format: str
    content: str
    metadata: Dict[str, Any]
    success: bool
    error: Optional[str] = None
    sections: Optional[List[Dict]] = None
    word_count: int = 0
    char_count: int = 0


class DocumentParser:
    """
    Universal document parser with support for multiple formats

    Usage:
        parser = DocumentParser()
        result = parser.parse('document.pdf')
        if result.success:
            print(result.content)
    """

    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.json': self._parse_json,
            '.jsonl': self._parse_jsonl,
            '.csv': self._parse_csv,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
        }

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse any supported document format

        Args:
            file_path: Path to document file

        Returns:
            ParsedDocument with content and metadata
        """
        file_path = Path(file_path)

        # Initialize result
        result = ParsedDocument(
            file_path=str(file_path),
            format=file_path.suffix.lower(),
            content="",
            metadata={
                'filename': file_path.name,
                'size_bytes': 0,
                'parsed_at': datetime.now().isoformat()
            },
            success=False
        )

        # Check file exists
        if not file_path.exists():
            result.error = f"File not found: {file_path}"
            return result

        # Get file size
        result.metadata['size_bytes'] = file_path.stat().st_size

        # Check format supported
        if result.format not in self.supported_formats:
            result.error = f"Unsupported format: {result.format}"
            return result

        # Parse with appropriate parser
        try:
            parser_func = self.supported_formats[result.format]
            parser_func(file_path, result)
            result.success = True

            # Calculate word/char counts
            result.word_count = len(result.content.split())
            result.char_count = len(result.content)

        except Exception as e:
            result.error = f"Parse error: {str(e)}"
            result.success = False

        return result

    def _parse_text(self, file_path: Path, result: ParsedDocument):
        """Parse plain text files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            result.content = f.read()

        result.metadata['encoding'] = 'utf-8'
        result.metadata['lines'] = result.content.count('\n') + 1

    def _parse_markdown(self, file_path: Path, result: ParsedDocument):
        """Parse Markdown files with section extraction"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        result.content = content

        # Extract sections by headers
        sections = []
        current_section = None

        for line in content.split('\n'):
            # Check for headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                # Save previous section
                if current_section:
                    sections.append(current_section)

                # Start new section
                level = len(header_match.group(1))
                title = header_match.group(2)
                current_section = {
                    'level': level,
                    'title': title,
                    'content': ''
                }
            elif current_section:
                current_section['content'] += line + '\n'

        # Save last section
        if current_section:
            sections.append(current_section)

        result.sections = sections
        result.metadata['section_count'] = len(sections)

    def _parse_json(self, file_path: Path, result: ParsedDocument):
        """Parse JSON files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Store raw JSON as content
        result.content = json.dumps(data, indent=2)

        # Store parsed data in metadata
        result.metadata['json_data'] = data
        result.metadata['json_type'] = type(data).__name__

        if isinstance(data, list):
            result.metadata['item_count'] = len(data)
        elif isinstance(data, dict):
            result.metadata['keys'] = list(data.keys())

    def _parse_jsonl(self, file_path: Path, result: ParsedDocument):
        """Parse JSONL (JSON Lines) files"""
        lines = []
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    lines.append(json.loads(line))

        # Store as formatted JSON
        result.content = json.dumps(lines, indent=2)

        # Store parsed data
        result.metadata['jsonl_data'] = lines
        result.metadata['line_count'] = len(lines)

        # Sample first item structure
        if lines:
            result.metadata['sample_keys'] = list(lines[0].keys()) if isinstance(lines[0], dict) else None

    def _parse_csv(self, file_path: Path, result: ParsedDocument):
        """Parse CSV files"""
        rows = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            headers = reader.fieldnames
            for row in reader:
                rows.append(row)

        # Convert to formatted text
        result.content = json.dumps(rows, indent=2)

        # Store metadata
        result.metadata['csv_data'] = rows
        result.metadata['headers'] = headers
        result.metadata['row_count'] = len(rows)
        result.metadata['column_count'] = len(headers) if headers else 0

    def _parse_pdf(self, file_path: Path, result: ParsedDocument):
        """Parse PDF files"""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                # Extract all text
                text_parts = []
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    text_parts.append(text)

                result.content = '\n\n'.join(text_parts)

                # Metadata
                result.metadata['page_count'] = len(reader.pages)

                # PDF metadata if available
                if reader.metadata:
                    result.metadata['pdf_metadata'] = {
                        'title': reader.metadata.get('/Title', 'Unknown'),
                        'author': reader.metadata.get('/Author', 'Unknown'),
                        'subject': reader.metadata.get('/Subject', 'Unknown'),
                    }

        except ImportError:
            result.error = "PyPDF2 not installed. Run: pip install PyPDF2"
            result.success = False
        except Exception as e:
            result.error = f"PDF parse error: {str(e)}"
            result.success = False

    def _parse_docx(self, file_path: Path, result: ParsedDocument):
        """Parse Word documents"""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract all paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            result.content = '\n\n'.join(paragraphs)

            # Metadata
            result.metadata['paragraph_count'] = len(paragraphs)

            # Extract tables if present
            if doc.tables:
                result.metadata['table_count'] = len(doc.tables)

        except ImportError:
            result.error = "python-docx not installed. Run: pip install python-docx"
            result.success = False
        except Exception as e:
            result.error = f"DOCX parse error: {str(e)}"
            result.success = False

    def _parse_html(self, file_path: Path, result: ParsedDocument):
        """Parse HTML files"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract text (remove scripts and styles)
            for script in soup(['script', 'style']):
                script.decompose()

            # Get text
            result.content = soup.get_text(separator='\n', strip=True)

            # Metadata
            result.metadata['title'] = soup.title.string if soup.title else 'No title'
            result.metadata['links'] = len(soup.find_all('a'))
            result.metadata['images'] = len(soup.find_all('img'))

        except ImportError:
            result.error = "beautifulsoup4 not installed. Run: pip install beautifulsoup4"
            result.success = False
        except Exception as e:
            result.error = f"HTML parse error: {str(e)}"
            result.success = False


class YouTubeTranscriptParser:
    """
    Specialized parser for YouTube transcripts
    Works with youtube-transcript-api
    """

    def __init__(self, api_version: str = "0.6.1"):
        """
        Initialize with specific API version

        Args:
            api_version: youtube-transcript-api version (default: 0.6.1)
        """
        self.api_version = api_version

    def parse_transcript(self, video_id: str, languages: List[str] = None) -> ParsedDocument:
        """
        Parse YouTube video transcript

        Args:
            video_id: YouTube video ID (11 chars)
            languages: List of language codes (default: ['en'])

        Returns:
            ParsedDocument with transcript content
        """
        if languages is None:
            languages = ['en']

        result = ParsedDocument(
            file_path=f"youtube://{video_id}",
            format='.transcript',
            content="",
            metadata={
                'video_id': video_id,
                'languages_requested': languages,
                'parsed_at': datetime.now().isoformat()
            },
            success=False
        )

        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            from youtube_transcript_api import TranscriptsDisabled, NoTranscriptFound

            # Get transcript
            transcript = YouTubeTranscriptApi.get_transcript(
                video_id,
                languages=languages
            )

            # Extract full text
            full_text = ' '.join([entry['text'] for entry in transcript])
            result.content = full_text

            # Metadata
            result.metadata['segment_count'] = len(transcript)
            result.metadata['raw_transcript'] = transcript
            result.metadata['language_detected'] = transcript[0].get('language', 'unknown') if transcript else None

            # Calculate stats
            result.word_count = len(full_text.split())
            result.char_count = len(full_text)

            result.success = True

        except TranscriptsDisabled:
            result.error = "Transcripts are disabled for this video"
        except NoTranscriptFound:
            result.error = f"No transcript found in languages: {languages}"
        except ImportError:
            result.error = "youtube-transcript-api not installed. Run: pip install youtube-transcript-api==0.6.1"
        except Exception as e:
            result.error = f"Transcript error: {str(e)}"

        return result

    def parse_transcript_with_timestamps(self, video_id: str) -> ParsedDocument:
        """
        Parse transcript with timestamp preservation

        Returns formatted transcript with timestamps
        """
        result = self.parse_transcript(video_id)

        if result.success and 'raw_transcript' in result.metadata:
            # Format with timestamps
            formatted = []
            for entry in result.metadata['raw_transcript']:
                timestamp = self._format_timestamp(entry['start'])
                text = entry['text']
                formatted.append(f"[{timestamp}] {text}")

            result.content = '\n'.join(formatted)
            result.metadata['formatted'] = True

        return result

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Convert seconds to HH:MM:SS format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"


class BatchDocumentParser:
    """
    Parse multiple documents in batch
    """

    def __init__(self):
        self.parser = DocumentParser()
        self.youtube_parser = YouTubeTranscriptParser()

    def parse_directory(self, directory: Union[str, Path],
                       file_pattern: str = "*.*") -> List[ParsedDocument]:
        """
        Parse all matching files in directory

        Args:
            directory: Directory path
            file_pattern: Glob pattern (default: *.*)

        Returns:
            List of ParsedDocument results
        """
        directory = Path(directory)
        results = []

        for file_path in directory.glob(file_pattern):
            if file_path.is_file():
                result = self.parser.parse(file_path)
                results.append(result)

        return results

    def parse_multiple_files(self, file_paths: List[Union[str, Path]]) -> List[ParsedDocument]:
        """
        Parse multiple specific files

        Args:
            file_paths: List of file paths

        Returns:
            List of ParsedDocument results
        """
        results = []
        for file_path in file_paths:
            result = self.parser.parse(file_path)
            results.append(result)

        return results

    def get_summary(self, results: List[ParsedDocument]) -> Dict:
        """
        Generate summary statistics for batch parse

        Args:
            results: List of ParsedDocument results

        Returns:
            Summary dictionary
        """
        total = len(results)
        successful = sum(1 for r in results if r.success)
        failed = total - successful

        total_words = sum(r.word_count for r in results if r.success)
        total_chars = sum(r.char_count for r in results if r.success)

        formats = {}
        for r in results:
            formats[r.format] = formats.get(r.format, 0) + 1

        return {
            'total_files': total,
            'successful': successful,
            'failed': failed,
            'success_rate': (successful / total * 100) if total > 0 else 0,
            'total_words': total_words,
            'total_chars': total_chars,
            'formats': formats,
            'errors': [r.error for r in results if r.error]
        }


# ============================================================================
# Convenience Functions
# ============================================================================

def parse_document(file_path: Union[str, Path]) -> ParsedDocument:
    """
    Quick parse single document

    Usage:
        result = parse_document('file.pdf')
        if result.success:
            print(result.content)
    """
    parser = DocumentParser()
    return parser.parse(file_path)


def parse_youtube_transcript(video_id: str) -> ParsedDocument:
    """
    Quick parse YouTube transcript

    Usage:
        result = parse_youtube_transcript('dQw4w9WgXcQ')
        if result.success:
            print(result.content)
    """
    parser = YouTubeTranscriptParser()
    return parser.parse_transcript(video_id)


def parse_directory(directory: Union[str, Path], pattern: str = "*.*") -> List[ParsedDocument]:
    """
    Quick parse all files in directory

    Usage:
        results = parse_directory('documents/', '*.pdf')
        for r in results:
            if r.success:
                print(f"{r.metadata['filename']}: {r.word_count} words")
    """
    batch_parser = BatchDocumentParser()
    return batch_parser.parse_directory(directory, pattern)


# ============================================================================
# Example Usage
# ============================================================================

if __name__ == "__main__":
    # Example 1: Parse single document
    print("=" * 60)
    print("Example 1: Parse Single Document")
    print("=" * 60)

    result = parse_document("example.txt")
    if result.success:
        print(f"✓ Parsed: {result.metadata['filename']}")
        print(f"  Format: {result.format}")
        print(f"  Words: {result.word_count}")
        print(f"  Chars: {result.char_count}")
        print(f"  Content preview: {result.content[:200]}...")
    else:
        print(f"✗ Error: {result.error}")

    print()

    # Example 2: Parse YouTube transcript
    print("=" * 60)
    print("Example 2: Parse YouTube Transcript")
    print("=" * 60)

    result = parse_youtube_transcript("dQw4w9WgXcQ")
    if result.success:
        print(f"✓ Video: {result.metadata['video_id']}")
        print(f"  Segments: {result.metadata['segment_count']}")
        print(f"  Words: {result.word_count}")
        print(f"  Preview: {result.content[:200]}...")
    else:
        print(f"✗ Error: {result.error}")

    print()

    # Example 3: Batch parse directory
    print("=" * 60)
    print("Example 3: Batch Parse Directory")
    print("=" * 60)

    batch_parser = BatchDocumentParser()
    results = batch_parser.parse_directory("documents/", "*.md")
    summary = batch_parser.get_summary(results)

    print(f"Total files: {summary['total_files']}")
    print(f"Successful: {summary['successful']}")
    print(f"Failed: {summary['failed']}")
    print(f"Success rate: {summary['success_rate']:.1f}%")
    print(f"Total words: {summary['total_words']:,}")
    print(f"Formats: {summary['formats']}")
